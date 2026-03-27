"""Tests for cover letter generator module.

Tests the actual CoverLetterGenerator functionality — no OpenAI mocks needed
since the module now accepts text directly and formats it as .docx.
"""

import pytest
from pathlib import Path

from docx import Document

from src.cover_letter_gen import CoverLetterGenerator


@pytest.fixture
def gen():
    """Create a CoverLetterGenerator with default config."""
    return CoverLetterGenerator()


@pytest.fixture
def gen_with_config():
    """Create a CoverLetterGenerator with custom config."""
    return CoverLetterGenerator(config={"some_key": "some_value"})


@pytest.fixture
def candidate_profile():
    """Sample candidate profile for testing."""
    return {
        "name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "555-123-4567",
    }


@pytest.fixture
def sample_text():
    """Sample cover letter text."""
    return (
        "Dear Hiring Manager,\n\n"
        "I am excited about the Compliance Manager position at Acme Corp.\n\n"
        "With my background in regulatory compliance and risk management, "
        "I believe I would be a strong fit.\n\n"
        "Sincerely,\nJane Doe"
    )


class TestCoverLetterGenInit:
    def test_init_with_config(self, gen_with_config):
        assert gen_with_config.config == {"some_key": "some_value"}

    def test_init_default_config(self, gen):
        assert gen.config == {}

    def test_init_none_config(self):
        g = CoverLetterGenerator(config=None)
        assert g.config == {}


class TestReadExample:
    def test_read_example_returns_text(self, tmp_path):
        """read_example should return text content from a .docx file."""
        doc = Document()
        doc.add_paragraph("Example cover letter with professional tone.")
        doc.add_paragraph("Second paragraph of the example.")
        example_path = tmp_path / "example.docx"
        doc.save(str(example_path))

        result = CoverLetterGenerator.read_example(example_path)
        assert isinstance(result, str)
        assert "professional tone" in result
        assert "Second paragraph" in result

    def test_read_example_missing_file_returns_empty(self, tmp_path):
        """read_example should return empty string for missing file."""
        result = CoverLetterGenerator.read_example(tmp_path / "nonexistent.docx")
        assert result == ""

    def test_read_example_skips_blank_paragraphs(self, tmp_path):
        """read_example should skip paragraphs that are only whitespace."""
        doc = Document()
        doc.add_paragraph("Content paragraph")
        doc.add_paragraph("")  # blank
        doc.add_paragraph("   ")  # whitespace only
        doc.add_paragraph("Another paragraph")
        example_path = tmp_path / "example.docx"
        doc.save(str(example_path))

        result = CoverLetterGenerator.read_example(example_path)
        lines = result.split("\n")
        assert len(lines) == 2
        assert "Content paragraph" in lines[0]
        assert "Another paragraph" in lines[1]


class TestSaveCoverLetter:
    def test_save_creates_valid_docx(self, tmp_path, sample_text):
        """save_cover_letter should create a valid .docx with expected content."""
        output_path = tmp_path / "cover_letter.docx"

        result = CoverLetterGenerator.save_cover_letter(
            text=sample_text,
            candidate_name="Jane Doe",
            candidate_email="jane@example.com",
            candidate_phone="555-123-4567",
            output_path=output_path,
        )
        assert result == output_path
        assert output_path.exists()

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Compliance Manager" in full_text
        assert "Jane Doe" in full_text
        assert "jane@example.com" in full_text
        assert "555-123-4567" in full_text

    def test_save_creates_parent_dirs(self, tmp_path, sample_text):
        """save_cover_letter should create parent directories if needed."""
        output_path = tmp_path / "subdir" / "deep" / "cover_letter.docx"

        result = CoverLetterGenerator.save_cover_letter(
            text=sample_text,
            candidate_name="Jane Doe",
            candidate_email="jane@example.com",
            candidate_phone="",
            output_path=output_path,
        )
        assert output_path.exists()

    def test_save_handles_empty_contact(self, tmp_path):
        """save_cover_letter should handle empty email and phone gracefully."""
        output_path = tmp_path / "cover.docx"
        CoverLetterGenerator.save_cover_letter(
            text="Short letter.",
            candidate_name="Name",
            candidate_email="",
            candidate_phone="",
            output_path=output_path,
        )
        assert output_path.exists()

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Name" in full_text

    def test_save_preserves_paragraphs(self, tmp_path):
        """Each non-empty line should become its own paragraph."""
        text = "Line one.\n\nLine two.\n\nLine three."
        output_path = tmp_path / "multi.docx"

        CoverLetterGenerator.save_cover_letter(
            text=text,
            candidate_name="Test",
            candidate_email="t@t.com",
            candidate_phone="",
            output_path=output_path,
        )

        doc = Document(str(output_path))
        content_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 3 content paragraphs + signature name + contact line
        assert any("Line one." in p for p in content_paragraphs)
        assert any("Line two." in p for p in content_paragraphs)
        assert any("Line three." in p for p in content_paragraphs)


class TestGenerateAndSave:
    def test_generate_and_save_creates_docx(self, gen, candidate_profile, sample_text, tmp_path):
        """generate_and_save should produce a .docx file."""
        output_dir = tmp_path / "output"

        result = gen.generate_and_save(
            text=sample_text,
            candidate_profile=candidate_profile,
            company="Acme Corp",
            role="Compliance Manager",
            output_dir=output_dir,
        )
        assert Path(result).exists()
        assert str(result).endswith(".docx")

        doc = Document(str(result))
        assert len(doc.paragraphs) >= 1

    def test_generate_and_save_filename_pattern(self, gen, candidate_profile, sample_text, tmp_path):
        """Output filename should contain 'Danna_Dobi_Cover_Letter' and role."""
        output_dir = tmp_path / "output"

        result = gen.generate_and_save(
            text=sample_text,
            candidate_profile=candidate_profile,
            company="Acme",
            role="Compliance Manager",
            output_dir=output_dir,
        )
        filename = Path(result).name
        assert "Danna_Dobi_Cover_Letter" in filename
        assert "Compliance" in filename
        assert filename.endswith(".docx")

    def test_generate_and_save_truncates_long_role(self, gen, candidate_profile, sample_text, tmp_path):
        """Long role names should be truncated in the filename."""
        output_dir = tmp_path / "output"

        long_role = "Senior Compliance Manager, Risk and Governance | Full Time — Remote"
        result = gen.generate_and_save(
            text=sample_text,
            candidate_profile=candidate_profile,
            company="BigCo",
            role=long_role,
            output_dir=output_dir,
        )
        filename = Path(result).name
        # The role portion should be truncated at the first delimiter
        assert "Danna_Dobi_Cover_Letter" in filename
        assert len(filename) < 200

    def test_generate_and_save_uses_candidate_info(self, gen, sample_text, tmp_path):
        """The .docx should include candidate name/email/phone from profile."""
        output_dir = tmp_path / "output"
        profile = {
            "name": "Test User",
            "email": "test@example.com",
            "phone": "999-888-7777",
        }

        result = gen.generate_and_save(
            text=sample_text,
            candidate_profile=profile,
            company="TestCo",
            role="Analyst",
            output_dir=output_dir,
        )

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test User" in full_text
        assert "test@example.com" in full_text

    def test_generate_and_save_creates_output_dir(self, gen, candidate_profile, sample_text, tmp_path):
        """generate_and_save should create the output directory if it doesn't exist."""
        output_dir = tmp_path / "nonexistent" / "nested"

        result = gen.generate_and_save(
            text=sample_text,
            candidate_profile=candidate_profile,
            company="Co",
            role="Role",
            output_dir=output_dir,
        )
        assert Path(result).exists()
        assert output_dir.exists()


from unittest.mock import patch, MagicMock


class TestGenerateTextAI:
    """Tests for AI-powered cover letter generation."""

    def test_generates_with_llm(self):
        """generate_text should use LLM when available."""
        with patch("src.cover_letter_gen.CoverLetterGenerator.read_example", return_value="tone ref"), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume text"), \
             patch("src.llm_client.generate_text", return_value="AI generated cover letter text.\n\nBest regards,") as mock_llm:
            # Clear cache before test
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result = CoverLetterGenerator.generate_text("Acme Corp", "Compliance Analyst", "Looking for AML expert")
            assert "AI generated cover letter" in result
            mock_llm.assert_called_once()

    def test_falls_back_to_template_on_llm_failure(self):
        """generate_text should fall back to template when LLM returns empty string."""
        with patch("src.cover_letter_gen.CoverLetterGenerator.read_example", return_value="tone ref"), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume text"), \
             patch("src.llm_client.generate_text", return_value=""):
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result = CoverLetterGenerator.generate_text("Acme Corp", "Compliance Analyst", "Looking for AML expert")
            # Should contain template-generated content (OCC, Best regards, etc.)
            assert "OCC" in result
            assert "Best regards," in result

    def test_falls_back_to_template_on_exception(self):
        """generate_text should fall back to template when LLM raises an exception."""
        with patch("src.llm_client.generate_text", side_effect=Exception("API error")), \
             patch("src.cover_letter_gen.CoverLetterGenerator.read_example", side_effect=Exception("read failed")):
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result = CoverLetterGenerator.generate_text("TestCo", "Risk Analyst", "Risk management role")
            # Should still produce a cover letter via template
            assert "Best regards," in result
            assert "TestCo" in result

    def test_cache_returns_same_result(self):
        """generate_text should cache results by company+title."""
        with patch("src.cover_letter_gen.CoverLetterGenerator.read_example", return_value="tone ref"), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume text"), \
             patch("src.llm_client.generate_text", return_value="Cached letter content.\n\nBest regards,") as mock_llm:
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result1 = CoverLetterGenerator.generate_text("CacheCo", "Analyst", "desc")
            result2 = CoverLetterGenerator.generate_text("CacheCo", "Analyst", "different desc")
            assert result1 == result2
            # LLM should only be called once
            mock_llm.assert_called_once()

    def test_cache_different_keys(self):
        """Different company+title combos should not share cache."""
        with patch("src.cover_letter_gen.CoverLetterGenerator.read_example", return_value="tone ref"), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume text"), \
             patch("src.llm_client.generate_text", side_effect=["Letter A", "Letter B"]):
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result1 = CoverLetterGenerator.generate_text("CompanyA", "Role1", "desc")
            result2 = CoverLetterGenerator.generate_text("CompanyB", "Role2", "desc")
            assert result1 != result2

    def test_deslop_applied_to_llm_output(self):
        """LLM output should be run through deslop clean_text."""
        with patch("src.cover_letter_gen.CoverLetterGenerator.read_example", return_value="tone ref"), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume text"), \
             patch("src.llm_client.generate_text", return_value="I'm thrilled to leverage my skills"):
            import src.cover_letter_gen as cl_mod
            cl_mod._cover_letter_cache.clear()

            result = CoverLetterGenerator.generate_text("DeSlopCo", "Role", "desc")
            # deslop should have cleaned these
            assert "thrilled" not in result
            assert "leverage" not in result


class TestGenerateTextTemplate:
    """Tests for the template-based fallback (existing behavior preserved)."""

    def test_template_includes_company_and_title(self):
        """Template should reference the company and cleaned title."""
        import src.cover_letter_gen as cl_mod
        result = CoverLetterGenerator._generate_template_text(
            "BigBank", "Senior Compliance Officer", "AML compliance role"
        )
        assert "BigBank" in result
        assert "Senior Compliance Officer" in result

    def test_template_aml_highlights(self):
        """Template should pick AML-related highlights from description."""
        import src.cover_letter_gen as cl_mod
        result = CoverLetterGenerator._generate_template_text(
            "FinCo", "Analyst", "We need someone with AML and BSA experience"
        )
        assert "BSA/AML" in result

    def test_template_default_highlights(self):
        """Template should use default highlights when no keywords match."""
        import src.cover_letter_gen as cl_mod
        result = CoverLetterGenerator._generate_template_text(
            "GenericCo", "Associate", "General office work"
        )
        assert "Best regards," in result

    def test_template_deduplicates_title(self):
        """Template should handle doubled titles like 'Risk Analyst Risk Analyst'."""
        import src.cover_letter_gen as cl_mod
        result = CoverLetterGenerator._generate_template_text(
            "Co", "Risk Analyst Risk Analyst", "desc"
        )
        # Should only appear once in the letter body
        assert "Risk Analyst Risk Analyst" not in result
        assert "Risk Analyst" in result
