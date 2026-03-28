"""Tests for the deslop (AI-text cleanup) module."""

import os
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock


class TestCleanText:
    """Tests for clean_text() function."""

    def test_replaces_em_dash(self):
        from src.deslop import clean_text
        assert clean_text("this — that") == "this, that"

    def test_replaces_en_dash(self):
        from src.deslop import clean_text
        assert clean_text("this – that") == "this, that"

    def test_replaces_excited(self):
        from src.deslop import clean_text
        assert clean_text("I'm excited to apply") == "I'd like to apply"

    def test_replaces_thrilled(self):
        from src.deslop import clean_text
        assert clean_text("I'm thrilled to share") == "I'm writing to share"

    def test_replaces_leverage(self):
        from src.deslop import clean_text
        assert clean_text("leverage my skills") == "use my skills"
        assert clean_text("leveraging data") == "using data"
        assert clean_text("Leverage the platform") == "Use the platform"

    def test_replaces_utilize(self):
        from src.deslop import clean_text
        assert clean_text("utilize resources") == "use resources"
        assert clean_text("utilization rate") == "use rate"

    def test_replaces_synergy(self):
        from src.deslop import clean_text
        assert clean_text("team synergy") == "team alignment"
        assert clean_text("find synergies") == "find connections"

    def test_removes_filler_phrases(self):
        from src.deslop import clean_text
        assert clean_text("Great question about this") == " about this"
        assert clean_text("I'd be happy to help") == " help"

    def test_no_change_for_clean_text(self):
        from src.deslop import clean_text
        text = "I have 9 years of compliance experience."
        assert clean_text(text) == text

    def test_multiple_replacements(self):
        from src.deslop import clean_text
        text = "I'm excited to leverage my skills — they create synergy"
        result = clean_text(text)
        assert "excited" not in result
        assert "leverage" not in result
        assert " — " not in result
        assert "synergy" not in result


class TestCleanDocx:
    """Tests for clean_docx() function."""

    def test_nonexistent_file_returns_path(self, tmp_path):
        from src.deslop import clean_docx
        fake_path = tmp_path / "nonexistent.docx"
        result = clean_docx(fake_path)
        assert result == fake_path

    def test_cleans_docx_file(self, tmp_path):
        from docx import Document
        from src.deslop import clean_docx

        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("I'm excited to leverage my experience.")
        doc.save(str(doc_path))

        clean_docx(doc_path)

        cleaned_doc = Document(str(doc_path))
        text = cleaned_doc.paragraphs[0].text
        assert "excited" not in text
        assert "leverage" not in text

    def test_no_change_for_clean_docx(self, tmp_path):
        from docx import Document
        from src.deslop import clean_docx

        doc_path = tmp_path / "clean.docx"
        doc = Document()
        original_text = "Professional compliance experience."
        doc.add_paragraph(original_text)
        doc.save(str(doc_path))

        clean_docx(doc_path)

        cleaned_doc = Document(str(doc_path))
        assert cleaned_doc.paragraphs[0].text == original_text


class TestCleanDirectory:
    """Tests for clean_directory() function."""

    def test_cleans_all_docx_in_directory(self, tmp_path):
        from docx import Document
        from src.deslop import clean_directory

        for name in ["a.docx", "b.docx"]:
            doc = Document()
            doc.add_paragraph("I'm excited to leverage skills.")
            doc.save(str(tmp_path / name))

        clean_directory(tmp_path)

        for name in ["a.docx", "b.docx"]:
            doc = Document(str(tmp_path / name))
            assert "excited" not in doc.paragraphs[0].text
