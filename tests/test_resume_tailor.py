"""Tests for resume tailor module.

Tests the actual ResumeTailor functionality — no OpenAI mocks needed
since the module now does surgical .docx replacement without API calls.
"""

import pytest
from pathlib import Path

from docx import Document

from src.resume_tailor import ResumeTailor


@pytest.fixture
def tailor():
    """Create a ResumeTailor with default config."""
    return ResumeTailor()


@pytest.fixture
def tailor_with_config():
    """Create a ResumeTailor with custom config."""
    return ResumeTailor(config={"some_key": "some_value"})


@pytest.fixture
def base_resume(tmp_path):
    """Create a realistic base resume .docx with sections."""
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | 555-123-4567")
    doc.add_paragraph("")

    doc.add_paragraph("Professional Summary")
    doc.add_paragraph(
        "Experienced compliance professional with 9+ years in "
        "regulatory oversight, risk management, and audit."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Core Competencies")
    doc.add_paragraph("• Regulatory Compliance")
    doc.add_paragraph("• Risk Assessment")
    doc.add_paragraph("• Internal Audit")
    doc.add_paragraph("• BSA/AML")
    doc.add_paragraph("")

    doc.add_paragraph("Professional Experience")
    doc.add_paragraph("OCC — Bank Examiner — 2016-2023")
    doc.add_paragraph("Led compliance examinations of national banks.")
    doc.add_paragraph("")

    doc.add_paragraph("Education")
    doc.add_paragraph("B.S. Finance — University of California")

    path = tmp_path / "base_resume.docx"
    doc.save(str(path))
    return path


class TestResumeTailorInit:
    def test_init_with_config(self, tailor_with_config):
        assert tailor_with_config.config == {"some_key": "some_value"}

    def test_init_default_config(self, tailor):
        assert tailor.config == {}

    def test_init_none_config(self):
        t = ResumeTailor(config=None)
        assert t.config == {}


class TestReadResumeText:
    def test_reads_content_correctly(self, base_resume):
        """read_resume_text should return all non-empty paragraph text."""
        text = ResumeTailor.read_resume_text(base_resume)
        assert isinstance(text, str)
        assert "Jane Doe" in text
        assert "Professional Summary" in text
        assert "Regulatory Compliance" in text
        assert "Bank Examiner" in text

    def test_reads_simple_docx(self, tmp_path):
        """read_resume_text should handle a simple single-paragraph docx."""
        doc = Document()
        doc.add_paragraph("Simple content.")
        path = tmp_path / "simple.docx"
        doc.save(str(path))

        text = ResumeTailor.read_resume_text(path)
        assert "Simple content." in text

    def test_skips_blank_paragraphs(self, tmp_path):
        """read_resume_text should skip blank/whitespace-only paragraphs."""
        doc = Document()
        doc.add_paragraph("Content A")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("Content B")
        path = tmp_path / "blanks.docx"
        doc.save(str(path))

        text = ResumeTailor.read_resume_text(path)
        lines = text.split("\n")
        assert len(lines) == 2
        assert "Content A" in lines[0]
        assert "Content B" in lines[1]


class TestCloneAndTailor:
    def test_replaces_summary(self, base_resume, tmp_path):
        """clone_and_tailor should replace the Professional Summary text."""
        output_path = tmp_path / "tailored.docx"
        new_summary = "Tailored summary for compliance role at Acme Corp."

        result = ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
            new_summary=new_summary,
        )
        assert result == output_path
        assert output_path.exists()

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert new_summary in full_text
        # Original summary should be gone
        assert "9+ years in regulatory oversight" not in full_text

    def test_replaces_competencies(self, base_resume, tmp_path):
        """clone_and_tailor should replace bullet-pointed competencies."""
        output_path = tmp_path / "tailored.docx"
        new_competencies = [
            "Anti-Money Laundering",
            "Fraud Detection",
            "SOX Compliance",
            "Vendor Risk Management",
        ]

        result = ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
            new_competencies=new_competencies,
        )
        assert output_path.exists()

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Anti-Money Laundering" in full_text
        assert "Fraud Detection" in full_text
        assert "SOX Compliance" in full_text
        assert "Vendor Risk Management" in full_text

    def test_replaces_both_summary_and_competencies(self, base_resume, tmp_path):
        """clone_and_tailor should replace both summary and competencies."""
        output_path = tmp_path / "both.docx"
        new_summary = "Tailored for fintech compliance."
        new_competencies = ["FinTech Reg", "Crypto Compliance"]

        ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
            new_summary=new_summary,
            new_competencies=new_competencies,
        )

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert new_summary in full_text
        assert "FinTech Reg" in full_text
        assert "Crypto Compliance" in full_text

    def test_preserves_other_sections(self, base_resume, tmp_path):
        """clone_and_tailor should NOT modify Professional Experience or Education."""
        output_path = tmp_path / "preserved.docx"

        ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
            new_summary="New summary text.",
        )

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Bank Examiner" in full_text
        assert "B.S. Finance" in full_text

    def test_no_changes_when_none_provided(self, base_resume, tmp_path):
        """clone_and_tailor with no new_summary or new_competencies should clone as-is."""
        output_path = tmp_path / "clone.docx"

        ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
        )
        assert output_path.exists()

        # Content should match original
        original_text = ResumeTailor.read_resume_text(base_resume)
        cloned_text = ResumeTailor.read_resume_text(output_path)
        assert original_text == cloned_text

    def test_creates_parent_dirs(self, base_resume, tmp_path):
        """clone_and_tailor should create parent directories if needed."""
        output_path = tmp_path / "sub" / "dir" / "tailored.docx"

        ResumeTailor.clone_and_tailor(
            base_path=base_resume,
            output_path=output_path,
            new_summary="Nested output test.",
        )
        assert output_path.exists()


class TestTailorAndSave:
    def test_produces_docx_file(self, tailor, base_resume, tmp_path):
        """tailor_and_save should produce a .docx file in the output directory."""
        output_dir = tmp_path / "output"

        result = tailor.tailor_and_save(
            base_resume_path=base_resume,
            new_summary="Tailored for compliance role.",
            new_competencies=["AML", "BSA", "Fraud"],
            company="Acme Corp",
            role="Compliance Manager",
            output_dir=output_dir,
        )
        assert Path(result).exists()
        assert str(result).endswith(".docx")

    def test_filename_pattern(self, tailor, base_resume, tmp_path):
        """Output filename should contain 'Danna_Dobi_Resume' and role."""
        output_dir = tmp_path / "output"

        result = tailor.tailor_and_save(
            base_resume_path=base_resume,
            new_summary="Summary",
            new_competencies=["Skill"],
            company="Co",
            role="Risk Analyst",
            output_dir=output_dir,
        )
        filename = Path(result).name
        assert "Danna_Dobi_Resume" in filename
        assert "Risk_Analyst" in filename
        assert filename.endswith(".docx")

    def test_content_is_tailored(self, tailor, base_resume, tmp_path):
        """The generated .docx should contain the tailored content."""
        output_dir = tmp_path / "output"

        result = tailor.tailor_and_save(
            base_resume_path=base_resume,
            new_summary="Expert in fintech compliance and crypto regulation.",
            new_competencies=["Crypto AML", "DeFi Risk"],
            company="CoinBase",
            role="Compliance Lead",
            output_dir=output_dir,
        )

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "fintech compliance" in full_text
        assert "Crypto AML" in full_text

    def test_creates_output_dir(self, tailor, base_resume, tmp_path):
        """tailor_and_save should create the output directory if it doesn't exist."""
        output_dir = tmp_path / "nonexistent" / "nested"

        result = tailor.tailor_and_save(
            base_resume_path=base_resume,
            new_summary="S",
            new_competencies=[],
            company="Co",
            role="Role",
            output_dir=output_dir,
        )
        assert Path(result).exists()
        assert output_dir.exists()


from unittest.mock import patch, MagicMock


class TestTailorSummaryAI:
    """Tests for AI-powered resume summary generation."""

    def test_generates_with_llm(self):
        """tailor_summary should use LLM when available."""
        with patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="Danna's resume text"), \
             patch("src.llm_client.generate_text", return_value="AI-generated professional summary for compliance role.") as mock_llm:
            result = ResumeTailor.tailor_summary("Acme Corp", "Compliance Analyst", "Looking for AML expert")
            assert "AI-generated professional summary" in result
            mock_llm.assert_called_once()

    def test_falls_back_to_template_on_empty_result(self):
        """tailor_summary should fall back to template when LLM returns empty."""
        with patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume"), \
             patch("src.llm_client.generate_text", return_value=""):
            result = ResumeTailor.tailor_summary("Acme Corp", "Compliance Analyst", "AML BSA compliance")
            # Template fallback should produce text with OCC mention
            assert "OCC" in result or "Comptroller" in result

    def test_falls_back_to_template_on_exception(self):
        """tailor_summary should fall back to template when LLM raises exception."""
        with patch("src.llm_client.generate_text", side_effect=Exception("API down")), \
             patch("src.resume_tailor.ResumeTailor.read_resume_text", side_effect=Exception("file error")):
            result = ResumeTailor.tailor_summary("TestCo", "Risk Analyst", "Risk management role with audit")
            # Should still produce a summary via template
            assert "Compliance professional" in result or "compliance" in result.lower()

    def test_deslop_applied_to_llm_output(self):
        """LLM output should be run through deslop clean_text."""
        with patch("src.resume_tailor.ResumeTailor.read_resume_text", return_value="resume"), \
             patch("src.llm_client.generate_text", return_value="I'm excited to leverage my extensive experience"):
            result = ResumeTailor.tailor_summary("DeSlopCo", "Role", "desc")
            assert "leverage" not in result


class TestTailorSummaryTemplate:
    """Tests for the template-based summary fallback (existing behavior preserved)."""

    def test_template_aml_theme(self):
        """Template should use AML/BSA theme for relevant JDs."""
        result = ResumeTailor._tailor_summary_template(
            "FinBank", "AML Analyst", "BSA AML anti-money laundering compliance"
        )
        assert "BSA/AML" in result

    def test_template_risk_theme(self):
        """Template should use risk theme for risk-focused JDs."""
        result = ResumeTailor._tailor_summary_template(
            "RiskCo", "Risk Manager", "risk assessment risk management enterprise risk"
        )
        assert "risk assessment" in result

    def test_template_fintech_theme(self):
        """Template should use fintech theme for fintech JDs."""
        result = ResumeTailor._tailor_summary_template(
            "FinTechCo", "Compliance Lead", "fintech digital banking regulatory"
        )
        assert "fintech" in result.lower()

    def test_template_default_theme(self):
        """Template should use default theme when no keywords match."""
        result = ResumeTailor._tailor_summary_template(
            "GenericCo", "Associate", "General office work"
        )
        assert "compliance" in result.lower()
        assert "OCC" in result or "Comptroller" in result
