"""Cover letter generation module.

Creates a clean .docx cover letter matching the template's tone and formatting.
Uses LLM generation with template fallback.
"""

import logging
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches

from src.utils import ensure_dir, sanitize_filename

logger = logging.getLogger(__name__)

# Module-level cache
_cover_letter_cache = {}


class CoverLetterGenerator:
    """Generates a .docx cover letter from provided text."""

    def __init__(self, config: dict[str, Any] | None = None):
        self.config = config or {}

    @staticmethod
    def _generate_template_text(company: str, title: str, description: str) -> str:
        """Original template-based cover letter generation (fallback)."""
        from datetime import datetime
        import re

        # Clean title: remove LinkedIn artifacts (doubled text, "with verification", etc.)
        title = title.split("\n")[0].strip()
        title = title.replace(" with verification", "")
        # Strip trailing punctuation
        title = title.rstrip(".,;:")
        # De-duplicate if title is repeated (e.g. "Risk Analyst Risk Analyst")
        words = title.split()
        if len(words) >= 2 and len(words) % 2 == 0:
            half = len(words) // 2
            if [w.lower() for w in words[:half]] == [w.lower() for w in words[half:]]:
                title = " ".join(words[:half])
        # Clean up pipe/dash suffixes (e.g. "Risk Analyst | $100/hr Remote")
        for sep in ["|", "—", " - $"]:
            if sep in title:
                title = title.split(sep)[0].strip()
        # Remove redundant trailing "Position"/"Role"/"Job" (avoids "...Position position at")
        title = re.sub(r'\s+(position|role|job)\s*$', '', title, flags=re.IGNORECASE)

        desc_lower = description.lower()

        highlights = []
        if any(w in desc_lower for w in ["aml", "bsa", "anti-money laundering"]):
            highlights.append("BSA/AML compliance and transaction monitoring")
        if any(w in desc_lower for w in ["risk assessment", "risk management"]):
            highlights.append("risk assessment and controls evaluation")
        if any(w in desc_lower for w in ["audit", "testing", "quality control"]):
            highlights.append("audit program management and compliance testing")
        if any(w in desc_lower for w in ["regulatory", "examination", "examiner"]):
            highlights.append("regulatory examination support and exam-readiness")
        if any(w in desc_lower for w in ["monitoring", "surveillance"]):
            highlights.append("compliance monitoring and validation")
        if any(w in desc_lower for w in ["fraud", "investigation"]):
            highlights.append("fraud detection and investigation support")
        if any(w in desc_lower for w in ["governance", "policy"]):
            highlights.append("governance reporting and policy management")
        if any(w in desc_lower for w in ["fintech", "digital banking"]):
            highlights.append("fintech and digital banking compliance")
        if any(w in desc_lower for w in ["remediation", "corrective action"]):
            highlights.append("remediation tracking and confirmation testing")
        if any(w in desc_lower for w in ["documentation", "reporting"]):
            highlights.append("findings documentation and regulatory reporting")

        if not highlights:
            highlights = ["compliance monitoring and validation", "regulatory exam support"]

        highlights = highlights[:3]

        today = datetime.now().strftime("%B %d, %Y")

        letter = f"""{today}

Dear Hiring Manager,

I'm writing to apply for the {title} position at {company}. With 9+ years of compliance experience, including four years as a federal bank examiner with the OCC and subsequent roles in fintech and banking, I bring a practical understanding of what regulators expect and how to build programs that hold up under scrutiny.

My background maps directly to what this role requires. At the OCC, I conducted safety-and-soundness and compliance examinations of nationally chartered banks, covering {highlights[0]}. That examiner perspective carries into every role since: at Cross River Bank, I led compliance and operational risk audits; at Prime Trust, I managed audit and compliance monitoring programs across fintech operations.

"""
        if len(highlights) > 1:
            letter += (
                f"What I'd bring to {company} is hands-on experience in {highlights[1]}, "
                f"combined with the ability to translate complex findings into clear, actionable write-ups. "
            )

        if len(highlights) > 2:
            letter += (
                f"I also have direct experience with {highlights[2]}, "
                f"which I understand is central to this role. "
            )

        letter += f"""I hold both a CFE and CAMS certification, and my work consistently focuses on making compliance programs practical, well-documented, and ready for whatever questions come next.

I'd welcome the chance to discuss how my background fits what {company} is building.

Thank you for your time.

Best regards,"""

        return letter

    @staticmethod
    def generate_text(company: str, title: str, description: str) -> str:
        """Generate tailored cover letter text using LLM with template fallback."""
        # 1. Check cache
        cache_key = f"{company}|{title}"
        if cache_key in _cover_letter_cache:
            logger.info("Cover letter cache hit: %s", cache_key)
            return _cover_letter_cache[cache_key]

        # 2. Try LLM
        try:
            from src.resume_tailor import ResumeTailor
            from src.llm_client import generate_text as llm_generate
            from src.deslop import clean_text

            resume_text = ResumeTailor.read_resume_text("templates/base_resume.docx")
            tone_reference = CoverLetterGenerator.read_example()

            system_prompt = f"""You write cover letters for Danna Dobi, a compliance professional applying for jobs.

Rules:
- Max 4 paragraphs, 250-300 words total
- First paragraph: state the role and reference something specific about the company from the job description
- Middle paragraphs: connect Danna's actual experience to what the JD asks for. Use specific examples from her background at OCC (federal bank examiner), Cross River Bank (compliance/operational risk audits), and Prime Trust (audit and compliance monitoring)
- Last paragraph: short, confident close. No begging or desperation
- Mention CFE and CAMS certifications naturally within sentences, not as a list
- Sign off with "Best regards," only (no name after it)
- Write like a competent professional, not an AI. No "I'm thrilled", "leverage", "I'm excited to bring my expertise", "passionate about", "unique opportunity", or similar slop
- Use the tone reference below as a style guide — match its directness and confidence level

TONE REFERENCE:
{tone_reference}"""

            user_prompt = f"""Write a cover letter for this position:

Company: {company}
Role: {title}

Job Description:
{description[:3000]}

Danna's Background:
{resume_text[:2000]}"""

            result = llm_generate(system_prompt, user_prompt, max_tokens=800, temperature=0.7)

            if result:
                result = clean_text(result)
                _cover_letter_cache[cache_key] = result
                logger.info("Generated AI cover letter for %s at %s", title, company)
                return result

        except Exception as e:
            logger.warning("LLM cover letter generation failed: %s", e)

        # 3. Fallback to template
        logger.info("Falling back to template cover letter for %s at %s", title, company)
        text = CoverLetterGenerator._generate_template_text(company, title, description)
        _cover_letter_cache[cache_key] = text
        return text

    @staticmethod
    def read_example(path: str | Path = "templates/base_cover_letter.docx") -> str:
        """Read the example cover letter for reference."""
        path = Path(path)
        if not path.exists():
            return ""
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    @staticmethod
    def save_cover_letter(text: str, candidate_name: str, candidate_email: str,
                           candidate_phone: str, output_path: str | Path) -> Path:
        """Save cover letter text as a professionally formatted .docx."""
        output_path = Path(output_path)
        ensure_dir(output_path.parent)

        doc = Document()

        # Match template formatting
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Split text into paragraphs and write
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
            p.space_after = Pt(6)

        # Signature block
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        p.add_run(candidate_name).bold = True

        contact = " | ".join(x for x in [candidate_email, candidate_phone] if x)
        if contact:
            doc.add_paragraph(contact)

        doc.save(str(output_path))
        logger.info("Saved cover letter to %s", output_path)
        return output_path

    def generate_and_save(self, text: str, candidate_profile: dict,
                           company: str, role: str,
                           output_dir: str | Path) -> Path:
        """Save a cover letter .docx."""
        output_dir = ensure_dir(output_dir)
        # Human-like filename: Danna_Dobi_Cover_Letter_ShortTitle.docx
        short_role = role.split(",")[0].split("|")[0].split("—")[0].split("/")[0].strip()[:30]
        filename = sanitize_filename(f"Danna_Dobi_Cover_Letter_{short_role}") + ".docx"
        output_path = output_dir / filename

        return self.save_cover_letter(
            text=text,
            candidate_name=candidate_profile.get("name", ""),
            candidate_email=candidate_profile.get("email", ""),
            candidate_phone=candidate_profile.get("phone", ""),
            output_path=output_path,
        )
