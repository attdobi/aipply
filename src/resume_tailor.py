"""Resume tailoring module.

Clones the original .docx and surgically replaces the Professional Summary
and reorders Core Competencies. Uses LLM generation with template fallback.
"""

import logging
import os
import shutil
from pathlib import Path
from typing import Any

from docx import Document

from src.utils import ensure_dir, sanitize_filename

logger = logging.getLogger(__name__)


class ResumeTailor:
    """Clones a .docx resume and replaces summary + competencies."""

    def __init__(self, config: dict[str, Any] | None = None):
        self.config = config or {}

    @staticmethod
    def read_resume_text(path: str | Path) -> str:
        """Read all text from a .docx resume."""
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    @staticmethod
    def clone_and_tailor(base_path: str | Path, output_path: str | Path,
                          new_summary: str | None = None,
                          new_competencies: list[str] | None = None) -> Path:
        """Clone .docx and surgically replace summary and/or competencies.

        Preserves ALL original formatting — fonts, sizes, bold, bullets, spacing.
        Only the text content of targeted paragraphs is changed.
        """
        output_path = Path(output_path)
        ensure_dir(output_path.parent)

        # Byte-for-byte copy first
        shutil.copy2(str(base_path), str(output_path))

        doc = Document(str(output_path))
        current_section = None
        comp_idx = 0

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            # Detect section headers
            if "Professional Summary" in text:
                current_section = "summary"
                continue
            elif "Core Competencies" in text:
                current_section = "competencies"
                comp_idx = 0
                continue
            elif any(h in text for h in ["Professional Experience", "Education", "Certifications"]):
                current_section = None
                continue

            # Replace summary
            if current_section == "summary" and new_summary:
                _replace_paragraph_text(p, new_summary)
                current_section = None

            # Replace competencies (preserve bullet format)
            elif current_section == "competencies" and new_competencies and text.startswith("•"):
                if comp_idx < len(new_competencies):
                    _replace_paragraph_text(p, f"• {new_competencies[comp_idx]}")
                    comp_idx += 1

        doc.save(str(output_path))
        logger.info("Saved tailored resume to %s", output_path)
        return output_path

    @staticmethod
    def _tailor_summary_template(company: str, title: str, description: str) -> str:
        """Original template-based summary generation (fallback)."""
        desc_lower = description.lower()

        themes = []
        if any(w in desc_lower for w in ["aml", "anti-money laundering", "bsa", "bank secrecy"]):
            themes.append("aml_bsa")
        if any(w in desc_lower for w in ["fraud", "investigation", "suspicious activity"]):
            themes.append("fraud")
        if any(w in desc_lower for w in ["risk assessment", "risk management", "enterprise risk"]):
            themes.append("risk")
        if any(w in desc_lower for w in ["audit", "internal audit", "testing"]):
            themes.append("audit")
        if any(w in desc_lower for w in ["fintech", "digital banking", "financial technology"]):
            themes.append("fintech")
        if any(w in desc_lower for w in ["regulatory", "examination", "regulator", "occ", "fdic", "cfpb"]):
            themes.append("regulatory")
        if any(w in desc_lower for w in ["compliance monitoring", "monitoring program"]):
            themes.append("monitoring")
        if any(w in desc_lower for w in ["governance", "policy", "controls"]):
            themes.append("governance")

        base = (
            "Compliance professional with 9+ years of experience, including over four years "
            "as a federal bank examiner with the Office of the Comptroller of the Currency (OCC). "
            "That background means walking into any compliance environment with a regulator's eye: "
            "knowing what examiners look for, how monitoring programs get evaluated, "
            "and what makes documentation hold up."
        )

        if "aml_bsa" in themes or "fraud" in themes:
            middle = (
                " Since leaving the OCC, hands-on roles in fintech and banking have focused on "
                "compliance monitoring, fraud risk assessment, and BSA/AML program oversight, "
                "including transaction monitoring, suspicious activity reporting, and regulatory "
                "exam preparation."
            )
        elif "risk" in themes:
            middle = (
                " Since leaving the OCC, hands-on roles in fintech and banking have centered on "
                "risk assessment, compliance testing, and controls evaluation, with direct experience "
                "identifying emerging risk patterns, documenting findings, and keeping governance "
                "artifacts exam-ready."
            )
        elif "fintech" in themes:
            middle = (
                " Since leaving the OCC, hands-on roles at fintech-bank partnerships and digital "
                "banking institutions have focused on compliance monitoring, audit program management, "
                "and keeping governance artifacts exam-ready in fast-moving regulatory environments."
            )
        elif "audit" in themes:
            middle = (
                " Since leaving the OCC, hands-on roles in fintech and banking have focused on "
                "audit program management, compliance testing, and remediation validation, "
                "with a track record of keeping documentation in exam-ready condition."
            )
        else:
            middle = (
                " Since leaving the OCC, hands-on roles in fintech and banking have focused on "
                "compliance monitoring, reviewing customer-facing communications and marketing materials, "
                "validating remediation, and keeping governance artifacts exam-ready."
            )

        closing = (
            " Known for taking full ownership of assigned work and translating complex findings "
            "into summaries leadership can use."
        )

        return base + middle + closing

    @staticmethod
    def tailor_summary(company: str, title: str, description: str) -> str:
        """Generate tailored professional summary using LLM with template fallback."""
        try:
            from src.llm_client import generate_text as llm_generate
            from src.deslop import clean_text

            resume_text = ResumeTailor.read_resume_text("templates/base_resume.docx")

            system_prompt = """You write professional summaries for Danna Dobi's resume, a compliance professional.

Rules:
- Write exactly 3-4 sentences
- Emphasize the aspects of her background most relevant to the specific job description
- Use ONLY facts from her actual resume — do not invent experience she doesn't have
- Mention specific employers (OCC, Cross River Bank, Prime Trust) where relevant
- No AI slop: no "passionate", "leverage", "thrilled", "excited", "dynamic"
- Direct, confident, professional tone"""

            user_prompt = f"""Write a professional summary for this job application:

Company: {company}
Role: {title}

Job Description:
{description[:3000]}

Danna's Actual Background:
{resume_text[:2000]}"""

            result = llm_generate(system_prompt, user_prompt, max_tokens=300, temperature=0.7)

            if result:
                result = clean_text(result)
                logger.info("Generated AI resume summary for %s at %s", title, company)
                return result

        except Exception as e:
            logger.warning("LLM resume summary generation failed: %s", e)

        logger.info("Falling back to template summary for %s at %s", title, company)
        return ResumeTailor._tailor_summary_template(company, title, description)

    @staticmethod
    def tailor_competencies(description: str) -> list[str]:
        """Reorder core competencies based on JD."""
        desc_lower = description.lower()

        all_comps = [
            "Compliance Monitoring & Validation Execution",
            "Customer-Facing Channel Reviews (Communications, Marketing, Sales Practices)",
            "Findings Documentation, Classification & Escalation",
            "Remediation Follow-Up Testing & Confirmation Reviews",
            "Issue Intake, Tracking & Management",
            "Audit-Ready Evidence Preparation",
            "Regulatory Exam Support",
            "Policy Lifecycle Management",
            "Internal Controls & Testing",
            "Process Improvement & Documentation Quality",
        ]

        extras = {
            "aml": "BSA/AML Compliance & Transaction Monitoring",
            "fraud": "Fraud Risk Assessment & Investigation Support",
            "risk assessment": "Risk Assessment & Controls Evaluation",
            "governance": "Corporate Governance & Regulatory Reporting",
            "fintech": "Fintech & Digital Banking Compliance",
            "data analysis": "Data Analysis & Trend Identification",
            "third party": "Third-Party Risk Oversight",
            "vendor": "Third-Party Risk Oversight",
        }

        selected = []
        used = set()

        for keyword, comp in extras.items():
            if keyword in desc_lower and comp not in used and len(selected) < 2:
                selected.append(comp)
                used.add(comp)

        scored = []
        for comp in all_comps:
            score = sum(1 for word in comp.lower().split() if len(word) > 3 and word in desc_lower)
            scored.append((score, comp))
        scored.sort(key=lambda x: -x[0])

        for _, comp in scored:
            if comp not in used and len(selected) < 10:
                selected.append(comp)
                used.add(comp)

        return selected[:10]

    def tailor_and_save(self, base_resume_path: str | Path,
                         new_summary: str, new_competencies: list[str],
                         company: str, role: str,
                         output_dir: str | Path) -> Path:
        """Clone template and apply tailored content."""
        output_dir = ensure_dir(output_dir)
        # Clean role: remove LinkedIn artifacts
        role = role.split("\n")[0].strip().replace(" with verification", "")
        words = role.split()
        if len(words) >= 2 and len(words) % 2 == 0:
            half = len(words) // 2
            if [w.lower() for w in words[:half]] == [w.lower() for w in words[half:]]:
                role = " ".join(words[:half])
        # Human-like filename: Danna_Dobi_Resume_ShortTitle.docx
        short_role = role.split(",")[0].split("|")[0].split("—")[0].split("/")[0].strip()[:30]
        filename = sanitize_filename(f"Danna_Dobi_Resume_{short_role}") + ".docx"
        output_path = output_dir / filename
        return self.clone_and_tailor(base_resume_path, output_path, new_summary, new_competencies)


def _replace_paragraph_text(paragraph, new_text: str):
    """Replace paragraph text while preserving formatting of the first run."""
    if not paragraph.runs:
        paragraph.text = new_text
        return
    first_run = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = new_text
